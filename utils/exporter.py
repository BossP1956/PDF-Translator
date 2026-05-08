import io
import requests
import re
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.oxml.ns import qn

def create_word_from_md(translated_lines):
    doc = Document()
    
    # 设置全局中文字体
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for line in translated_lines:
        line = line.strip()
        if not line: continue
            
        # --- 1. 处理图片 (例如: ![alt](url)) ---
        img_match = re.search(r'!\[.*?\]\((.*?)\)', line)
        if img_match:
            img_url = img_match.group(1)
            try:
                img_data = requests.get(img_url, timeout=5).content
                img_stream = io.BytesIO(img_data)
                doc.add_picture(img_stream, width=Inches(5.0)) # 自动设为5英寸宽
                continue 
            except:
                p = doc.add_paragraph("[图片加载失败]")
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
                continue

        # --- 2. 处理标题 ---
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            
        # --- 3. 处理列表 ---
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            
        # --- 4. 处理表格符号 (保留原文) ---
        elif '|' in line and '--' not in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.name = 'Courier New' # 表格用等宽字体尝试对齐

        # --- 5. 处理普通正文 (含公式) ---
        else:
            doc.add_paragraph(line)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
